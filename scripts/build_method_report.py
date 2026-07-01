"""Build the final Beacon method report as DOCX and PDF.

The generated document is submission-oriented: it explains the current pipeline,
maps methods to reference papers, and keeps detailed artifacts in appendices.
"""

from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage, ImageDraw, ImageFont

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    LongTable,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "reports"
DOCX_PATH = OUT_DIR / "beacon_method_report_submission_repo_paths.docx"
PDF_PATH = OUT_DIR / "beacon_method_report_submission_repo_paths.pdf"
EXTRACTED_IMAGE_DIR = OUT_DIR / "extracted_images"
CURRENT_PIPELINE_DIAGRAM = EXTRACTED_IMAGE_DIR / "current_pipeline.png"
METHOD_LAYER_DIAGRAM = EXTRACTED_IMAGE_DIR / "method_layers.png"
WALKTHROUGH_DIAGRAM = EXTRACTED_IMAGE_DIR / "example_walkthrough.png"
RETRY_LOOP_DIAGRAM = EXTRACTED_IMAGE_DIR / "sql_attempt_loop.png"
TEST_RESULTS_DIR = ROOT / "tests" / "test_results"
TEST_RESULTS_JSON = TEST_RESULTS_DIR / "master_plan_evaluation_results.json"
TEST_RESULTS_HTML = TEST_RESULTS_DIR / "report.html"
TEST_RESULTS_REPO_PATH = "tests/test_results/"


INK = RGBColor(0x1F, 0x29, 0x37)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = "F2F4F7"
BORDER = "D9DEE7"


IMPORTANT_TERMS = [
    "Beacon",
    "metadata-grounded",
    "semantic profiles",
    "value evidence",
    "schema linking",
    "read-only",
    "retry",
    "BIRD-dev",
    "latest 10-question rerun",
    "55% accuracy",
]

PATH_EXTENSIONS = (".py", ".json", ".md", ".html", ".docx", ".pdf", ".png", ".csv", ".toml")
BOOKMARK_ID = 1


PAPERS = [
    {
        "id": "P1",
        "title": "Automatic Metadata Extraction for Text-to-SQL",
        "area": "Semantic metadata",
        "benchmark": "BIRD dev/minidev and BIRD test",
        "result": "67.41% BIRD test without oracle hints; 77.14 with oracle hints",
        "beacon_use": "Motivates compact semantic profiles, sample rows, and metadata-first schema context.",
    },
    {
        "id": "P2",
        "title": "E-SQL: Direct Schema Linking via Question Enrichment",
        "area": "Question enrichment and value grounding",
        "benchmark": "BIRD development and test",
        "result": "66.29% EX on BIRD test; 56.45% EX on BIRD dev with DeepSeek Coder 7B Instruct 1.5",
        "beacon_use": "Motivates matched evidence blocks and explicit value hints such as payment aliases.",
    },
    {
        "id": "P3",
        "title": "SOMA-SQL: Resolving Multi-Source Ambiguity in NL-to-SQL via Synthetic Log and Execution Probing",
        "area": "Ambiguity and probing",
        "benchmark": "Spider 2.0-Lite and BIRD",
        "result": "72.02% EX on Spider 2.0-Lite; average +13.0% EX over non-interactive baselines",
        "beacon_use": "Motivates future safe probes for ambiguous values and columns.",
    },
    {
        "id": "P4",
        "title": "The Death of Schema Linking? Text-to-SQL in the Age of Well-Reasoned Language Models",
        "area": "Recall-first schema context",
        "benchmark": "BIRD",
        "result": "71.83% BIRD accuracy, first at the time discussed by the paper",
        "beacon_use": "Motivates avoiding aggressive hard pruning when missing a required table is costly.",
    },
    {
        "id": "P5",
        "title": "KaSLA: Knapsack Optimization-based Schema Linking",
        "area": "Budgeted schema linking",
        "benchmark": "BIRD-dev and Spider-dev",
        "result": "CodeS-15B BIRD-dev improved from 57.63% EX to 60.17% EX; table Recall+ 99.15",
        "beacon_use": "Motivates ranked schema selection and planned core/support/suppressed table roles.",
    },
    {
        "id": "P6",
        "title": "LitE-SQL: A Lightweight and Efficient Text-to-SQL Framework with Vector-based Schema Linking and Execution-Guided Self-Correction",
        "area": "Vector linking and execution correction",
        "benchmark": "BIRD dev and Spider 1.0",
        "result": "68.45% EX on BIRD-dev; 88.45% EX on Spider 1.0",
        "beacon_use": "Motivates local vector schema retrieval and feeding execution errors into retry.",
    },
    {
        "id": "P7",
        "title": "Improving Retrieval-augmented Text-to-SQL with AST-based Ranking and Schema Pruning",
        "area": "Retrieval ranking and schema pruning",
        "benchmark": "Spider dev, Spider variants, CSpider",
        "result": "86.6% EX / 77.3% EM on Spider dev with GPT-4 + ASTRES + Graphix-T5",
        "beacon_use": "Motivates future SQL-shape-aware ranking and output discipline.",
    },
    {
        "id": "P8",
        "title": "Arctic-Text2SQL-R1: Simple Rewards, Strong Reasoning in Text-to-SQL",
        "area": "Execution correctness",
        "benchmark": "BIRD dev/test and Spider test",
        "result": "7B model: 68.9% EX on BIRD-dev, 68.5% EX on BIRD-test, 88.8% EX on Spider-test",
        "beacon_use": "Motivates treating execution feedback as the main correction signal.",
    },
    {
        "id": "P9",
        "title": "DAIL-SQL: Text-to-SQL Empowered by Large Language Models",
        "area": "Prompting and examples",
        "benchmark": "Spider, Spider-Realistic, BIRD",
        "result": "86.6% EX on Spider; outperformed DIN-SQL on BIRD dev/test",
        "beacon_use": "Motivates stable prompt section ordering and structurally relevant examples.",
    },
    {
        "id": "P10",
        "title": "RoboPhD: Self-Improving Text-to-SQL Through Autonomous Agent Evolution",
        "area": "Agent evolution and memory",
        "benchmark": "BIRD, DS-1000, ARC-AGI",
        "result": "73.67% EX on BIRD test and 71.30% EX on BIRD dev with evolved agents",
        "beacon_use": "Motivates feedback-example candidates as a lightweight future memory mechanism.",
    },
    {
        "id": "P11",
        "title": "CycleSQL: Data-Based Self-Explanations for Text-to-SQL",
        "area": "Result review",
        "benchmark": "Spider validation/test and Spider robustness variants",
        "result": "RESDSQL reached 82.0% translation accuracy on Spider validation (+2.6) and 81.6% on Spider test (+3.2)",
        "beacon_use": "Motivates reviewing result samples before composing the final answer.",
    },
    {
        "id": "P12",
        "title": "Memo-SQL: Structured Decomposition and Experience-Driven Self-Correction",
        "area": "Decomposition and self-correction",
        "benchmark": "BIRD dev, BIRD dev-new, Spider dev, CHESS-SDS",
        "result": "67.6% EX on BIRD dev, 68.5% EX on BIRD dev-new, 86.5% EX on Spider dev",
        "beacon_use": "Motivates decomposition, retry messages, and later experience-driven example promotion.",
    },
    {
        "id": "P13",
        "title": "Reliable Text-to-SQL with Adaptive Abstention",
        "area": "Abstention and schema reliability",
        "benchmark": "BIRD, Spider-dev, Spider-test",
        "result": "DeepSeek-7B with RTS schema: 64.72% EX on BIRD and 88.90% on Spider-dev",
        "beacon_use": "Motivates readable failure when schema coverage is insufficient.",
    },
    {
        "id": "P14",
        "title": "TrustSQL: Benchmarking Text-to-SQL Reliability with Penalty-Based Scoring",
        "area": "Reliability evaluation",
        "benchmark": "TrustSQL from ATIS, Advising, EHRSQL, Spider",
        "result": "Lenient reliability scores up to 75.0 on Spider; stricter penalties expose unsafe over-answering",
        "beacon_use": "Motivates read-only validation, failure states, and future abstention metrics.",
    },
]


METHOD_MAP = [
    (
        "Offline semantic profile enrichment",
        "Profile CSVs, sample rows, min/max/mean, distinct counts, top values",
        "P1, P2",
        "Implemented in src/beacon/indexing/profiles.py",
    ),
    (
        "Generic question signals",
        "Detect metrics, filters, dates, limits, time grain, weak lexical schema matches",
        "P2, P9",
        "Implemented in src/beacon/linking/question_signals.py",
    ),
    (
        "Metadata and value grounding",
        "Map user terms to table.column=value evidence with aliases, scores, confidence, and pinning",
        "P1, P2, P3",
        "Implemented in src/beacon/linking/metadata_grounding.py; probing remains future work",
    ),
    (
        "Vector schema retrieval",
        "Search table/column schema records with sentence-transformer embeddings or deterministic hash fallback",
        "P5, P6, P7",
        "Implemented in src/beacon/linking/schema_index.py and vector_store.py",
    ),
    (
        "Schema graph expansion",
        "Add bounded join paths and bridge join columns after selecting tables",
        "P4, P5, P7",
        "Implemented in src/beacon/linking/schema_graph.py and schema_linking.py",
    ),
    (
        "Structural few-shot example retrieval",
        "Rank examples by table, metric, filter, time grain, and question overlap",
        "P9, P12",
        "Implemented in src/beacon/linking/example_retrieval.py",
    ),
    (
        "Prompt assembly",
        "Order SQL rules, matched evidence, join paths, schema docs, examples, question, SQL marker",
        "P2, P9",
        "Implemented in src/beacon/runtime/prompting.py",
    ),
    (
        "SQL safety validation",
        "Allow one read-only SELECT/WITH statement, block writes/admin functions, enforce selected schema context",
        "P13, P14",
        "Implemented in src/beacon/runtime/sql.py",
    ),
    (
        "Read-only execution",
        "Run validated SQL in PostgreSQL read-only repeatable-read transaction with timeout and preview limit",
        "P8, P11, P14",
        "Implemented in src/beacon/runtime/sql.py",
    ),
    (
        "Reviewer-guided retry",
        "Review status, error, SQL, and result summary; retry up to three attempts with SQL/value/retrieval repair",
        "P6, P8, P11, P12",
        "Implemented in src/beacon/runtime/pipeline.py and retry.py",
    ),
    (
        "Feedback example candidates",
        "Optionally save accepted attempts for later curated few-shot promotion",
        "P10, P12",
        "Implemented in src/beacon/runtime/feedback_examples.py behind BEACON_SAVE_EXAMPLE_CANDIDATES",
    ),
    (
        "Planned benchmark extension",
        "Schema focus roles and output projection discipline for BIRD-style exact scoring",
        "P4, P5, P7, P9, P14",
        "Planned in sibling benchmark notes; not yet in current basic-mvp implementation",
    ),
]


CONTENTS_ENTRIES = [
    ("1. Executive Summary", "Thesis, scope, strongest method idea, and current limitation.", "sec_exec", 3),
    ("2. System Context", "Dataset, semantic model, evidence artifacts, and schema scope.", "sec_context", 4),
    ("3. Current Pipeline", "Runtime and offline flow from UI/CLI to answer.", "sec_pipeline", 5),
    ("4. Related Work Positioning", "Why these papers matter and how Beacon differs from full systems.", "sec_related", 6),
    ("5. Research-to-Method Traceability", "Compact mapping from method step to reference papers and implementation status.", "sec_traceability", 7),
    ("6. Method Details", "Layer-by-layer method, outputs, implementation file, and failure behavior.", "sec_method", 8),
    ("7. Concrete Example Walkthrough", "Complex question traced through every layer.", "sec_walkthrough", 11),
    ("8. Implementation Call Flow", "Function-level call contracts.", "sec_call_flow", 13),
    ("9. Evaluation Status", "Current focused tests, latest 10-question run, and BIRD-dev result.", "sec_evaluation", 14),
    ("10. Limitations and Next Work", "Prioritized risks and next actions.", "sec_limitations", 15),
    ("Appendix A. Reference Paper Matrix", "Reference paper matrix.", "sec_appendix_a", 16),
    ("Appendix B. Commands and Artifacts", "Commands, artifacts, and wrapper files.", "sec_appendix_b", 17),
    ("Appendix C. Horizontal Diagram Recreation Prompt", "Diagram prompt for future redraws.", "sec_appendix_c", 17),
]


EVALUATION_ROWS = [
    (
        "Focused tests",
        "47 passing in the current notes",
        "Retrieval, semantic profiles, prompts, SQL validation, retry, feedback examples, schema graph, smoke behavior",
    ),
    (
        "Local index build",
        "Successful",
        "Semantic JSON was enriched and local schema vector artifacts were built under data/indices/local_vectors/.",
    ),
    (
        "Latest 10-question rerun",
        "English: 7 completed, 3 API/model-channel errors; Vietnamese: 6 completed, 2 failed, 2 API/model-channel errors",
        "Run on 2026-06-30 with uv run python tests/test_cases/run_master_plan_tests.py. Results are stored in tests/test_results/.",
    ),
    (
        "Latest result artifacts",
        "JSON and HTML report generated",
        "tests/test_results/master_plan_evaluation_results.json and tests/test_results/report.html record the latest run.",
    ),
    (
        "Benchmark planning",
        "Best achieved BIRD-dev result: 55% accuracy",
        "This is the best current BIRD-dev result from the benchmark planning track; schema focus and projection discipline remain the next improvement targets.",
    ),
]


LIMITATION_ROWS = [
    (
        "High",
        "Schema selection is recall-first and can include extra tables.",
        "Add schema focus roles such as core, support, and suppressed before benchmark runs.",
    ),
    (
        "High",
        "Reviewer acceptance is useful but not a proof of semantic correctness.",
        "Compare alternative SQL shapes or add stronger result probes for high-risk questions.",
    ),
    (
        "Medium",
        "Value grounding uses profiles and aliases, but does not safely probe ambiguous values yet.",
        "Add controlled execution probing for ambiguous high-value filters.",
    ),
    (
        "Medium",
        "SQL validation is readable but regex-based.",
        "Keep the readable validator for the MVP; consider a parser if nested SQL edge cases become common.",
    ),
    (
        "Medium",
        "Accepted feedback examples are saved as candidates, not automatically promoted.",
        "Review candidates manually before they enter curated few-shot examples.",
    ),
    (
        "Low",
        "Vietnamese and multilingual question support are not systematic.",
        "Treat multilingual support as a later product extension after the English method stabilizes.",
    ),
]


FUNCTION_MAP = [
    ("UI entry", "beacon.app.ui.handle_question", "Calls beacon.runtime.pipeline.ask_database for Gradio."),
    ("CLI entry", "beacon.runtime.pipeline.main", "Reads a question and prints the final answer."),
    ("Question orchestration", "answer_question", "Splits obviously independent questions and runs sections."),
    ("Section orchestration", "answer_section", "Retrieves context, builds prompt, loops SQL attempts, returns answer."),
    ("Retrieval entry", "linking.retrieval.retrieve_context", "Loads metadata/examples and calls link_schema."),
    ("Schema linking", "linking.schema_linking.link_schema", "Combines signals, grounding, vectors, fallback, joins, examples."),
    ("Prompt", "runtime.prompting.build_sql_prompt", "Builds the SQL-only prompt context."),
    ("SQL attempt", "runtime.pipeline.run_sql_attempt", "Requests SQL, validates, executes, reviews."),
    ("Retry", "runtime.retry.classify_retry_need", "Chooses SQL retry, value repair, or retrieval repair."),
    ("SQL validation", "runtime.sql.validate_sql", "Enforces read-only single-statement selected-schema SQL."),
    ("Execution", "runtime.sql.run_query", "Runs read-only PostgreSQL transaction with count and preview."),
    ("Final answer", "pipeline_tools.compose_final_answer", "Composes natural-language answer from accepted attempt/result."),
]


METHOD_DEEP_DETAILS = [
    {
        "layer": "Offline semantic profiling",
        "input": "data/semantic_model/*.json plus data/processed/*.csv",
        "mechanic": "For each table, Beacon streams the CSV, accumulates per-column profile state, and finalizes JSON-safe profiles. Text columns keep sample and top values; date columns keep min/max; numeric columns keep min/max and mean when useful.",
        "output": "Enriched semantic model with sample_rows and column.profile fields, plus prompt-ready schema docs.",
        "papers": "P1, P2",
        "code": "src/beacon/indexing/profiles.py",
        "failure": "If a CSV is missing, the static semantic JSON still loads, but value/profile grounding becomes weaker.",
    },
    {
        "layer": "Question signal extraction",
        "input": "Original user question and optional semantic model text",
        "mechanic": "The extractor normalizes text, removes stopwords, detects intent phrases, metric phrases, filter phrases, years, numeric limits, capitalized entities, and weak lexical table/column overlaps.",
        "output": "signals dict: terms, entities, dates, numbers, intents, metrics, filters, time_grain, weak_tables, weak_columns, reasons.",
        "papers": "P2, P9",
        "code": "src/beacon/linking/question_signals.py",
        "failure": "Signals are deliberately weak. If they miss a schema item, vector hits, metadata grounding, and fallback rules can still recover it.",
    },
    {
        "layer": "Metadata and value grounding",
        "input": "Question text plus enriched semantic model values and aliases",
        "mechanic": "Beacon builds grounding candidates from profile values, top values, and manual aliases. It scores candidates, computes margin, labels confidence, and pins only strong unambiguous evidence.",
        "output": "Evidence rows such as term -> table.column = SQL literal with source, score, confidence, status, and pin flag.",
        "papers": "P1, P2, P3",
        "code": "src/beacon/linking/metadata_grounding.py",
        "failure": "Ambiguous terms are shown to the LLM but not forced. Execution probing from SOMA-SQL is a future improvement.",
    },
    {
        "layer": "Vector schema retrieval",
        "input": "Question embedding and schema records built from table/column metadata",
        "mechanic": "The vector store saves records.json, vectors.npy, and manifest.json. At runtime, Beacon searches table and column records by cosine similarity. Tests can use deterministic hash embeddings.",
        "output": "Ranked vector hits with record metadata and score.",
        "papers": "P5, P6, P7",
        "code": "src/beacon/linking/schema_index.py and vector_store.py",
        "failure": "Hash embeddings are deterministic but crude. Production sentence-transformer embeddings should reduce noisy extra tables.",
    },
    {
        "layer": "Hybrid schema linking",
        "input": "Signals, grounded evidence, vector hits, compatibility fallback, semantic relations, few-shot examples",
        "mechanic": "link_schema merges all schema signals into selected tables/columns, expands bounded join paths, builds schema docs, assesses coverage, and ranks examples.",
        "output": "linked_context: selected_tables, selected_columns, join_paths, schema_docs, example_docs, coverage, evidence, fallback_needs.",
        "papers": "P4, P5, P7, P9",
        "code": "src/beacon/linking/schema_linking.py",
        "failure": "The current MVP is recall-first and can include extra prompt tables. Schema focus roles are planned benchmark work.",
    },
    {
        "layer": "Prompt assembly",
        "input": "Linked context and original question",
        "mechanic": "The prompt is assembled in a fixed order: SQL rules, matched evidence, join paths, relevant schema docs, example queries, question, and final SQL marker.",
        "output": "One SQL-only prompt passed into the section message history.",
        "papers": "P2, P9",
        "code": "src/beacon/runtime/prompting.py",
        "failure": "If schema context is too broad, the prompt remains valid but the model may choose a distracting table or column.",
    },
    {
        "layer": "SQL validation and execution",
        "input": "Raw model SQL and allowed selected tables",
        "mechanic": "Beacon strips fences, requires one SELECT/WITH statement, blocks write/admin keywords and unsafe functions, extracts referenced tables, then executes inside a read-only repeatable-read PostgreSQL transaction.",
        "output": "Result dict with columns, preview rows, and total count, or a validation/execution error.",
        "papers": "P8, P11, P13, P14",
        "code": "src/beacon/runtime/sql.py",
        "failure": "Regex validation is readable but not a full SQL parser. Complex SQL edge cases remain a known limitation.",
    },
    {
        "layer": "Review and retry",
        "input": "Attempt status, SQL, error, result summary, original question, message history",
        "mechanic": "The reviewer returns strict JSON. If not satisfied, retry classification chooses SQL retry, value repair, or retrieval repair. Retrieval repair can add a known missing table and join path to context.",
        "output": "Accepted attempt and final answer, or a failed section with attempt history.",
        "papers": "P6, P8, P11, P12",
        "code": "src/beacon/runtime/pipeline.py, pipeline_tools.py, retry.py",
        "failure": "A plausible but semantically wrong result can still pass review. Future work should add stronger probes and alternative-query comparison.",
    },
]


EXAMPLE_WALKTHROUGH_ROWS = [
    (
        "Question",
        "For paid Apple Pay orders in the East region during 2022, what are the top 5 product categories by net revenue, discount rate, and profit margin?",
        "This one question combines value aliases, profile values, a year filter, top-k ranking, grouped product analytics, geography, and revenue/profit formulas.",
    ),
    (
        "Question signals",
        "intents={aggregation, top_n, group_by}; metrics={profit, discount, revenue}; filters={payment_filter, product_filter, date_filter}; date=2022; limit=5; entities={Apple Pay, East}.",
        "The question is routed toward a filtered, ranked aggregate over order, item, product, and geography evidence.",
    ),
    (
        "Metadata grounding",
        "\"apple pay\" -> orders.payment_method='apple_pay'; \"paid\" -> orders.order_status='paid'; \"east\" -> geography.region='east'. All three are high-confidence pinned evidence.",
        "The prompt receives concrete SQL literals instead of asking the model to guess spelling or casing.",
    ),
    (
        "Vector and fallback selection",
        "Fallback and weak signals contribute order_date, category, quantity, unit_price, discount_amount, cogs, and region. Hash-vector hits also pull some inventory/sales context.",
        "The essential columns are retrieved. Extra recall-first context remains visible as a current precision limitation.",
    ),
    (
        "Selected tables",
        "customers, geography, inventory, order_items, orders, products, sales.",
        "orders, order_items, products, and geography are essential. customers, inventory, and sales are extra in this offline hash-embedding inspection.",
    ),
    (
        "Join paths",
        "order_items.order_id -> orders.order_id; order_items.product_id -> products.product_id; geography.zip -> orders.zip; plus bridge paths through customers/inventory.",
        "The necessary path for this answer is orders -> order_items -> products and orders -> geography.",
    ),
    (
        "Retrieved examples",
        "Top examples include revenue by product segment for the current year and city generated the most revenue.",
        "Together they demonstrate revenue formula + date filtering and geography joins.",
    ),
    (
        "Prompt output",
        "SQL rules + join paths + schema docs + example queries + original question + SQL marker.",
        "The LLM is asked for SQL only, no markdown or visible chain-of-thought.",
    ),
    (
        "Expected SQL shape",
        "Join orders, order_items, products, and geography; filter paid/apple_pay/east/2022; group by p.category; compute net revenue, discount rate, and profit margin; order by net revenue; limit 5.",
        "This target shape shows why the pipeline needs value grounding, join paths, examples, and SQL validation in one run.",
    ),
]


EXPECTED_WALKTHROUGH_SQL = (
    "SELECT p.category,\n"
    "       SUM(oi.quantity * oi.unit_price - oi.discount_amount) AS net_revenue,\n"
    "       SUM(oi.discount_amount) / NULLIF(SUM(oi.quantity * oi.unit_price), 0) AS discount_rate,\n"
    "       SUM(oi.quantity * oi.unit_price - oi.discount_amount - p.cogs * oi.quantity)\n"
    "         / NULLIF(SUM(oi.quantity * oi.unit_price - oi.discount_amount), 0) AS profit_margin\n"
    "FROM orders o\n"
    "JOIN order_items oi ON oi.order_id = o.order_id\n"
    "JOIN products p ON p.product_id = oi.product_id\n"
    "JOIN geography g ON g.zip = o.zip\n"
    "WHERE o.order_status = 'paid'\n"
    "  AND o.payment_method = 'apple_pay'\n"
    "  AND g.region = 'east'\n"
    "  AND o.order_date >= DATE '2022-01-01'\n"
    "  AND o.order_date < DATE '2023-01-01'\n"
    "GROUP BY p.category\n"
    "ORDER BY net_revenue DESC\n"
    "LIMIT 5"
)


DIAGRAM_RECREATE_PROMPT = (
    "Create a clean horizontal architecture diagram for the Beacon NL-to-SQL MVP. "
    "Use a left-to-right flow on a white background with readable labels. Stages: "
    "1) App entry: Gradio UI in beacon.app.ui or CLI in beacon.runtime.pipeline; "
    "2) answer_question / answer_section: split obvious independent questions and manage attempts; "
    "3) retrieve_context: load semantic JSON and few-shot examples; "
    "4) link_schema: question_signals, metadata_grounding, vector schema hits, compatibility fallback, schema_graph joins, example_retrieval; "
    "5) build_sql_prompt: SQL rules, matched evidence, join paths, schema docs, examples, question; "
    "6) request_sql and call_llm; "
    "7) runtime.sql: clean_sql, validate_sql, run_query in read-only PostgreSQL; "
    "8) review_attempt and retry: SQL retry, value repair, or retrieval repair, max 3 attempts; "
    "9) compose_final_answer and optional feedback example candidate. "
    "Show offline indexing above the retrieval stage: profiles.py and builder.py enrich semantic_model JSON and local vector files. "
    "Use blue for app/runtime, green for linking/indexing, amber for LLM/prompt, and red for validation/retry."
)


def load_font(size: int, bold: bool = False):
    """Load a readable local font, falling back to PIL's default."""
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    """Wrap text to a pixel width for generated diagrams."""
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: str,
    outline: str = "#CBD5E1",
    title_size: int = 22,
    body_size: int = 17,
) -> None:
    """Draw one rounded diagram box."""
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    title_font = load_font(title_size, bold=True)
    body_font = load_font(body_size)
    draw.text((x1 + 18, y1 + 14), title, fill="#0F172A", font=title_font)
    y = y1 + title_size + 28
    for part in body.splitlines():
        for line in wrap_text(draw, part, body_font, x2 - x1 - 36):
            draw.text((x1 + 18, y), line, fill="#334155", font=body_font)
            y += body_size + 8


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    """Draw a simple arrow between boxes."""
    draw.line([start, end], fill="#64748B", width=3)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        points = [(ex, ey), (ex - 12, ey - 7), (ex - 12, ey + 7)]
    else:
        points = [(ex, ey), (ex + 12, ey - 7), (ex + 12, ey + 7)]
    draw.polygon(points, fill="#64748B")


def create_current_pipeline_diagram() -> None:
    """Generate a horizontal diagram that matches the current package layout."""
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    img = PILImage.new("RGB", (1500, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.text((50, 36), "Beacon Current Pipeline", fill="#1F2937", font=load_font(34, bold=True))
    draw.text(
        (50, 82),
        "Current package flow: app entrypoints -> runtime orchestration -> linking/retrieval -> SQL validation/execution -> review/retry -> answer.",
        fill="#475569",
        font=load_font(18),
    )

    top_boxes = [
        ((50, 170, 240, 330), "App Entry", "Gradio UI\nCLI", "#EEF6FF"),
        ((300, 170, 490, 330), "Runtime", "split sections\nrun attempts", "#EEF2FF"),
        ((550, 170, 740, 330), "Retrieve", "semantic JSON\nfew-shot examples", "#F0FDF4"),
        ((800, 170, 990, 330), "Prompt + LLM", "rules/evidence\nschema/examples", "#FFF7ED"),
        ((1050, 170, 1240, 330), "SQL Runtime", "validate SQL\nread-only run", "#F8FAFC"),
        ((1300, 170, 1450, 330), "Answer", "review/retry\nfinal answer", "#F0FDFA"),
    ]
    for xy, title, body, fill in top_boxes:
        draw_box(draw, xy, title, body, fill, title_size=24, body_size=22)
    for start_x, end_x in [(240, 300), (490, 550), (740, 800), (990, 1050), (1240, 1300)]:
        draw_arrow(draw, (start_x, 250), (end_x, 250))

    draw.text((50, 410), "Inside link_schema", fill="#1F2937", font=load_font(26, bold=True))
    link_boxes = [
        ((50, 470, 245, 625), "Signals", "metrics, filters\ndates, limits", "#ECFDF5"),
        ((295, 470, 490, 625), "Grounding", "aliases + values\npinned evidence", "#ECFDF5"),
        ((540, 470, 735, 625), "Vector Hits", "schema records\nvector search", "#ECFDF5"),
        ((785, 470, 980, 625), "Fallback", "compatibility\nschema hints", "#F7FEE7"),
        ((1030, 470, 1225, 625), "Schema Graph", "join paths\nbridge columns", "#ECFDF5"),
        ((1275, 470, 1470, 625), "Examples", "ranked by shape\nand overlap", "#ECFDF5"),
    ]
    for xy, title, body, fill in link_boxes:
        draw_box(draw, xy, title, body, fill, title_size=24, body_size=21)
    for start_x, end_x in [(245, 295), (490, 540), (735, 785), (980, 1030), (1225, 1275)]:
        draw_arrow(draw, (start_x, 548), (end_x, 548))
    draw_arrow(draw, (645, 330), (645, 470))
    draw_arrow(draw, (1370, 470), (895, 330))

    offline_boxes = [
        ((310, 760, 600, 900), "Offline Indexing", "profiles.py enriches JSON\nbuilder.py writes vectors", "#F8FAFC"),
        ((690, 760, 980, 900), "Persistent Artifacts", "semantic_model/*.json\nlocal_vectors/*\nfew_shot_queries.json", "#F8FAFC"),
    ]
    for xy, title, body, fill in offline_boxes:
        draw_box(draw, xy, title, body, fill, title_size=24, body_size=20)
    draw_arrow(draw, (600, 830), (690, 830))
    draw_arrow(draw, (835, 760), (645, 625))

    draw.text(
        (50, 930),
        "Legend: blue=runtime, green=linking, amber=LLM prompt, gray=data and SQL safety.",
        fill="#475569",
        font=load_font(18, bold=True),
    )
    img.save(CURRENT_PIPELINE_DIAGRAM)


def create_method_layer_diagram() -> None:
    """Generate a data/artifact flow visualization for the method section."""
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    img = PILImage.new("RGB", (1500, 950), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 42), "Beacon Method Artifact Flow", fill="#1F2937", font=load_font(34, bold=True))
    draw.text(
        (60, 90),
        "What changes shape as a question moves through the method.",
        fill="#475569",
        font=load_font(19),
    )
    boxes = [
        ((60, 170, 305, 330), "Source Evidence", "CSV rows\nsemantic JSON\nfew-shot SQL", "#F8FAFC"),
        ((380, 170, 625, 330), "Enriched Metadata", "profiles\nsample rows\nschema docs", "#EEF6FF"),
        ((700, 170, 945, 330), "Retrieval Records", "table/column text\nlocal vectors\nexample signals", "#F0FDF4"),
        ((1020, 170, 1265, 330), "Linked Context", "evidence\nselected schema\njoin paths", "#ECFDF5"),
    ]
    for xy, title, body, fill in boxes:
        draw_box(draw, xy, title, body, fill, title_size=24, body_size=21)
    for start_x, end_x in [(305, 380), (625, 700), (945, 1020)]:
        draw_arrow(draw, (start_x, 250), (end_x, 250))

    lower = [
        ((220, 520, 465, 690), "Prompt Package", "rules\nevidence\nschema/examples\nquestion", "#FFF7ED"),
        ((615, 520, 860, 690), "SQL Attempt", "clean SQL\nvalidation status\nexecution result", "#F8FAFC"),
        ((1010, 520, 1255, 690), "Review State", "accepted answer\nretry instruction\nfeedback candidate", "#FEF2F2"),
    ]
    for xy, title, body, fill in lower:
        draw_box(draw, xy, title, body, fill, title_size=24, body_size=21)
    draw_arrow(draw, (1140, 330), (345, 520))
    draw_arrow(draw, (465, 605), (615, 605))
    draw_arrow(draw, (860, 605), (1010, 605))
    draw_arrow(draw, (1110, 520), (760, 330))

    draw.text(
        (60, 830),
        "Thesis: the system improves by preserving grounded evidence, then adding precision through joins, validation, review, and targeted retry.",
        fill="#475569",
        font=load_font(19, bold=True),
    )
    img.save(METHOD_LAYER_DIAGRAM)


def create_retry_loop_diagram() -> None:
    """Generate the SQL attempt and retry loop visualization."""
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    img = PILImage.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "SQL Attempt Loop", fill="#1F2937", font=load_font(34, bold=True))
    boxes = [
        ((60, 140, 360, 280), "Request SQL", "Model returns SQL only", "#EEF6FF"),
        ((480, 140, 780, 280), "Validate", "Read-only, one statement, safe functions, selected tables only", "#F8FAFC"),
        ((900, 140, 1200, 280), "Execute", "Read-only transaction, count, preview rows", "#F0FDF4"),
        ((900, 430, 1200, 570), "Review", "Strict JSON: satisfied, reason, retry instructions", "#FFF7ED"),
        ((480, 430, 780, 570), "Repair", "SQL retry, value repair, or retrieval repair", "#FEF2F2"),
        ((60, 430, 360, 570), "Answer", "Accepted result becomes final answer", "#F0FDFA"),
    ]
    for box in boxes:
        draw_box(draw, *box)
    draw_arrow(draw, (360, 210), (480, 210))
    draw_arrow(draw, (780, 210), (900, 210))
    draw_arrow(draw, (1050, 280), (1050, 430))
    draw_arrow(draw, (900, 500), (780, 500))
    draw_arrow(draw, (480, 500), (360, 500))
    draw_arrow(draw, (630, 430), (210, 280))
    draw.text((540, 615), "Loop limit: MAX_SQL_ATTEMPTS = 3", fill="#475569", font=load_font(18, bold=True))
    img.save(RETRY_LOOP_DIAGRAM)


def create_walkthrough_diagram() -> None:
    """Generate the concrete example walkthrough visualization."""
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    img = PILImage.new("RGB", (1300, 1600), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "Example Walkthrough: Paid Apple Pay Orders, East Region, 2022", fill="#1F2937", font=load_font(30, bold=True))
    boxes = [
        ((190, 135, 1110, 255), "1. Question", "Top 5 product categories by net revenue, discount rate, and profit margin for paid Apple Pay orders in East during 2022.", "#EEF6FF"),
        ((190, 315, 1110, 435), "2. Signals", "aggregation + group_by + top_n; metrics=profit, discount, revenue; date=2022; limit=5", "#F0FDF4"),
        ((190, 495, 1110, 615), "3. Grounding", "Pinned evidence: apple_pay, paid, east. These become explicit SQL literals.", "#F8FAFC"),
        ((190, 675, 1110, 795), "4. Schema Selection", "Essential tables present: orders, order_items, products, geography. Extra recall-first tables are noted.", "#FFF7ED"),
        ((190, 855, 1110, 975), "5. Join Paths", "orders -> order_items -> products; orders.zip -> geography.zip", "#F0FDFA"),
        ((190, 1035, 1110, 1155), "6. Retrieved Examples", "Revenue-by-segment/year and city-revenue examples shape the SQL pattern.", "#F5F3FF"),
        ((190, 1215, 1110, 1335), "7. Prompt Package", "Rules + evidence + joins + schema docs + examples + question + SQL marker.", "#EEF2FF"),
        ((190, 1395, 1110, 1515), "8. Target SQL Shape", "Filter paid/apple_pay/east/2022; group by category; compute net revenue, discount rate, profit margin; limit 5.", "#ECFDF5"),
    ]
    for xy, title, body, fill in boxes:
        draw_box(draw, xy, title, body, fill, title_size=23, body_size=19)
    for y in [255, 435, 615, 795, 975, 1155, 1335]:
        draw_arrow(draw, (650, y), (650, y + 60))
    draw.text(
        (190, 1545),
        "Observation: grounding succeeds; extra tables reveal the next schema-focus improvement.",
        fill="#475569",
        font=load_font(18, bold=True),
    )
    img.save(WALKTHROUGH_DIAGRAM)


def create_visual_assets() -> None:
    """Validate externally recreated visual assets used by the report."""
    missing = [
        path
        for path in [
            CURRENT_PIPELINE_DIAGRAM,
            METHOD_LAYER_DIAGRAM,
            RETRY_LOOP_DIAGRAM,
            WALKTHROUGH_DIAGRAM,
        ]
        if not path.exists()
    ]
    if missing:
        raise FileNotFoundError("Missing recreated report diagram(s): " + ", ".join(str(path) for path in missing))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if bold:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        if color is not None:
            run.font.color.rgb = color
    else:
        add_formatted_runs(p, text, size=9)


def is_path_token(token: str) -> bool:
    stripped = token.strip(".,;:()[]'\"")
    return (
        "/" in stripped
        or "\\" in stripped
        or stripped.endswith(PATH_EXTENSIONS)
        or stripped.startswith(("src", "docs", "data", "tests", "scripts", "uv.lock", "pyproject"))
    )


def split_formatted_tokens(text: str) -> list[tuple[str, str]]:
    important = "|".join(re.escape(term) for term in sorted(IMPORTANT_TERMS, key=len, reverse=True))
    pattern = re.compile(rf"({important})|(\S+)", re.IGNORECASE)
    parts: list[tuple[str, str]] = []
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            parts.append(("normal", text[pos : match.start()]))
        token = match.group(0)
        if match.group(1):
            parts.append(("important", token))
        elif is_path_token(token):
            parts.append(("path", token))
        else:
            parts.append(("normal", token))
        pos = match.end()
    if pos < len(text):
        parts.append(("normal", text[pos:]))
    return parts


def add_formatted_runs(paragraph, text: str, size: float = 11) -> None:
    for kind, value in split_formatted_tokens(text):
        run = paragraph.add_run(value)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        if kind == "important":
            run.bold = True
            run.font.color.rgb = INK
        elif kind == "path":
            run.italic = True
            run.font.color.rgb = MUTED


def add_docx_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    add_formatted_runs(p, text)
    return p


def add_bookmark(paragraph, name: str) -> None:
    global BOOKMARK_ID
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(BOOKMARK_ID))
    bookmark_start.set(qn("w:name"), name)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(BOOKMARK_ID))
    BOOKMARK_ID += 1
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)


def add_bookmarked_heading(doc: Document, text: str, bookmark: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    add_bookmark(p, bookmark)
    return p


def add_internal_hyperlink(paragraph, text: str, anchor: str, page: int | None = None) -> None:
    label = text if page is None else f"{text}    {page}"
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = label
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_external_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)


def repeat_table_header(row) -> None:
    """Mark a Word table row as a repeating header."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_docx_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple[str, ...]],
    widths: list[float],
    caption: str | None = None,
) -> None:
    if caption:
        add_table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    table.autofit = False
    repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        table.rows[0].cells[index].width = Inches(widths[index])
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=INK)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].width = Inches(widths[index])
            set_cell_text(cells[index], value)
    doc.add_paragraph()


def add_docx_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8FAFC")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for line_index, line in enumerate(code.splitlines()):
        if line_index:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def setup_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def add_page_furniture(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Beacon Method Report"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header.runs:
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Polished method draft - 2026-06-30"
    if footer.runs:
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = MUTED


def add_docx_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Beacon Method and Pipeline Report")
    run.font.name = "Calibri"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Submission-ready technical method report")
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    for label, value in [
        ("Project", "Beacon / basic-mvp"),
        ("Report date", "2026-06-30"),
        ("Primary template spine", "Oxford Engineering style adapted for a technical method report"),
        ("Scope", "Current MVP implementation plus related benchmark-planning notes from the sibling Beacon workspace"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(label + ": ")
        r1.bold = True
        r1.font.color.rgb = INK
        p.add_run(value)
    doc.add_paragraph()


def add_docx_contents(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Contents", level=1)
    add_docx_paragraph(
        doc,
        "The entries below use static page numbers from the generated PDF and hyperlink directly to the corresponding section in the DOCX.",
    )
    for title, purpose, bookmark, page in CONTENTS_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, title, bookmark, page)
        note = doc.add_paragraph()
        note.paragraph_format.left_indent = Inches(0.25)
        note.paragraph_format.space_after = Pt(4)
        run = note.add_run(purpose)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def traceability_rows() -> list[tuple[str, str, str]]:
    """Return compact traceability rows with citation-style references."""
    return [
        (step, f"{mechanic} [{papers}].", status)
        for step, mechanic, papers, status in METHOD_MAP
    ]


def method_detail_paragraphs(detail: dict) -> list[str]:
    """Write one method layer in prose with inline reference IDs."""
    return [
        (
            f"Input: {detail['input']}. {detail['mechanic']} This follows the grounding, retrieval, prompting, or reliability pattern most relevant to this layer [{detail['papers']}]."
        ),
        (
            f"Output: {detail['output']} Implementation: {detail['code']}. Failure behavior: {detail['failure']}"
        ),
    ]


def add_docx_report() -> None:
    doc = Document()
    setup_docx_styles(doc)
    add_page_furniture(doc)
    add_docx_title(doc)
    add_docx_contents(doc)

    doc.add_page_break()
    add_bookmarked_heading(doc, "1. Executive Summary", "sec_exec", level=1)
    add_docx_paragraph(
        doc,
        "Beacon is a metadata-grounded natural-language-to-SQL MVP. Its current core pipeline is: "
        "profile data, retrieve grounded schema context, build a SQL-only prompt, generate one read-only "
        "PostgreSQL query, validate and execute it safely, review the result, retry when useful, and then "
        "compose the final answer."
    )
    add_docx_paragraph(
        doc,
        "The method is intentionally simple: plain dictionaries, small helper modules, and focused tests. The strongest current design idea is not model cleverness but grounding: semantic profiles, value evidence, local vector retrieval, join paths, and examples.",
    )
    add_docx_paragraph(
        doc,
        "The main current limitation is that selected schema is still recall-first and can include distracting context. The sibling benchmark plan therefore adds schema focus and output discipline as future work. Reference papers are mapped step-by-step so the reader can see which part of the system came from which research idea.",
    )

    doc.add_page_break()
    add_bookmarked_heading(doc, "2. System Context", "sec_context", level=1)
    add_docx_paragraph(
        doc,
        "Beacon currently works over a local e-commerce PostgreSQL schema with seven tables: customers, orders, order_items, products, geography, sales, and inventory. "
        "Each table has a semantic JSON file under data/semantic_model/. The offline indexing layer enriches those files with compact data evidence instead of dumping raw data into prompts."
    )
    add_docx_paragraph(
        doc,
        "The semantic layer stores table-level evidence such as semantic name, description, grain, relations, and sample rows. It also stores column-level evidence such as names, types, descriptions, profiles, aliases, and value aliases where useful.",
    )
    add_docx_paragraph(
        doc,
        "Profile fields include null count, distinct count, sample values, numeric/date bounds, means for non-ID numeric fields, and top categorical values. Few-shot examples live in data/few_shot_queries.json and are enriched with structural signals during indexing.",
    )

    doc.add_page_break()
    add_bookmarked_heading(doc, "3. Current Pipeline", "sec_pipeline", level=1)
    add_docx_paragraph(
        doc,
        "The current basic-mvp implementation is organized in four layers: indexing, linking, runtime, and app. "
        "The important source files are under src/beacon/, with compatibility wrappers kept at the package root."
    )
    if CURRENT_PIPELINE_DIAGRAM.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(CURRENT_PIPELINE_DIAGRAM), width=Inches(6.5))
        add_caption(doc, "Figure 1. Current horizontal Beacon pipeline, generated from the package layout.")
    add_docx_paragraph(
        doc,
        "At runtime, the UI or CLI passes the question into answer_question and answer_section. The section runner retrieves schema context, assembles a SQL-only prompt, requests SQL from the configured model, validates and executes the query through runtime.sql, reviews the result, and either retries with targeted repair or composes the final answer."
    )

    doc.add_page_break()
    add_bookmarked_heading(doc, "4. Related Work Positioning", "sec_related", level=1)
    add_docx_paragraph(
        doc,
        "The reference set was chosen around four practical questions: how to enrich schema context, how to avoid losing required tables, how to organize prompts and examples, and how to recover from SQL errors. Beacon is not trying to reproduce a full competition system such as DAIL-SQL or DIN-SQL. Instead, it extracts the pieces that fit a small, inspectable MVP: semantic metadata, evidence-first linking, stable prompt order, read-only validation, and a bounded retry loop."
    )
    add_docx_paragraph(
        doc,
        "The strongest thesis is that Beacon improves usefulness through grounding rather than model cleverness. Papers on metadata extraction and question enrichment motivate profile values and matched evidence [P1, P2]. Schema-linking and retrieval papers motivate recall-preserving schema selection with later precision work [P4-P7]. Prompting, retry, and reliability papers motivate the fixed prompt layout, execution feedback, result review, and readable failure states [P8-P14]."
    )
    add_docx_paragraph(
        doc,
        "This positioning also explains the current limitation: the MVP prefers to retrieve enough context, even when that includes extra tables. The benchmark-oriented next step is not a bigger orchestration framework, but a sharper schema-focus layer that preserves recall while reducing distracting columns and projection mismatches."
    )

    doc.add_page_break()
    add_bookmarked_heading(doc, "5. Research-to-Method Traceability", "sec_traceability", level=1)
    add_docx_paragraph(
        doc,
        "This table is the core traceability view. It distinguishes implemented MVP mechanics from planned benchmark extensions."
    )
    add_docx_table(
        doc,
        ["Pipeline step", "Method source", "Implementation status"],
        traceability_rows(),
        [1.35, 3.0, 2.0],
        caption="Table 2. Research-to-method traceability matrix.",
    )

    doc.add_page_break()
    add_bookmarked_heading(doc, "6. Method Details", "sec_method", level=1)
    add_docx_paragraph(
        doc,
        "This section describes each layer as a method, not just as a file. Reference IDs are cited inline in the same style as a paper; Appendix A expands the titles and benchmark notes."
    )
    if METHOD_LAYER_DIAGRAM.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(METHOD_LAYER_DIAGRAM), width=Inches(6.3))
        add_caption(doc, "Figure 2. Data and artifact flow through Beacon's method layers.")

    for detail in METHOD_DEEP_DETAILS:
        doc.add_heading(detail["layer"], level=2)
        for paragraph in method_detail_paragraphs(detail):
            add_docx_paragraph(doc, paragraph)

        if detail["layer"] == "SQL validation and execution" and RETRY_LOOP_DIAGRAM.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(RETRY_LOOP_DIAGRAM), width=Inches(6.3))
            add_caption(doc, "Figure 3. SQL validation, execution, review, and retry loop.")

    doc.add_page_break()
    add_bookmarked_heading(doc, "7. Concrete Example Walkthrough", "sec_walkthrough", level=1)
    add_docx_paragraph(
        doc,
        "The walkthrough below uses the question 'For paid Apple Pay orders in the East region during 2022, what are the top 5 product categories by net revenue, discount rate, and profit margin?' and records the actual retrieval/linking outputs from a deterministic hash-embedding inspection run. This example is intentionally more demanding: it needs value grounding, date recognition, multi-table joins, grouped metrics, top-k ranking, and a precise SQL output shape."
    )
    if WALKTHROUGH_DIAGRAM.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(WALKTHROUGH_DIAGRAM), width=Inches(6.3))
        add_caption(doc, "Figure 4. Layer-by-layer walkthrough for the complex paid-Apple-Pay regional revenue question.")
    add_docx_table(
        doc,
        ["Layer", "Observed output", "Interpretation"],
        EXAMPLE_WALKTHROUGH_ROWS,
        [1.15, 2.8, 2.45],
        caption="Table 3. Observed outputs at each walkthrough layer.",
    )
    doc.add_heading("Expected SQL Shape For The Example", level=2)
    add_docx_code_block(doc, EXPECTED_WALKTHROUGH_SQL)

    doc.add_page_break()
    add_bookmarked_heading(doc, "8. Implementation Call Flow", "sec_call_flow", level=1)
    add_docx_paragraph(
        doc,
        "This table is intentionally function-level: it shows the call contract after the method layers have already explained the research mechanics."
    )
    add_docx_table(
        doc,
        ["Role", "Function or module", "Call contract"],
        FUNCTION_MAP,
        [1.25, 2.2, 2.95],
        caption="Table 4. Function-level call flow.",
    )

    doc.add_page_break()
    add_bookmarked_heading(doc, "9. Evaluation Status", "sec_evaluation", level=1)
    add_docx_paragraph(
        doc,
        "Evaluation is still MVP-scale, but it is concrete enough to show the pipeline is testable. The strongest evidence is the combination of focused tests, successful index construction, the latest 10-question rerun, and the benchmark planning note that the best achieved BIRD-dev result is 55% accuracy."
    )
    p = doc.add_paragraph()
    add_formatted_runs(p, "Latest run artifacts folder: ")
    add_external_hyperlink(p, TEST_RESULTS_REPO_PATH, TEST_RESULTS_REPO_PATH)
    add_docx_table(
        doc,
        ["Evidence", "Current status", "Notes"],
        EVALUATION_ROWS,
        [1.35, 1.45, 3.65],
        caption="Table 5. Current evaluation evidence and caveats.",
    )

    doc.add_page_break()
    add_bookmarked_heading(doc, "10. Limitations and Next Work", "sec_limitations", level=1)
    add_docx_paragraph(
        doc,
        "The next work should follow the same thesis as the method: keep the pipeline grounded and inspectable, then improve precision where the walkthrough exposes extra context."
    )
    add_docx_table(
        doc,
        ["Priority", "Limitation", "Next work"],
        LIMITATION_ROWS,
        [0.7, 2.65, 3.05],
        caption="Table 6. Prioritized limitations and next actions.",
    )

    doc.add_page_break()
    add_bookmarked_heading(doc, "Appendix A. Reference Paper Matrix", "sec_appendix_a", level=1)
    add_docx_table(
        doc,
        ["ID", "Paper", "Area", "Benchmark/result", "Beacon use"],
        [(p["id"], p["title"], p["area"], p["benchmark"] + "; " + p["result"], p["beacon_use"]) for p in PAPERS],
        [0.45, 1.8, 1.0, 1.8, 1.35],
        caption="Table A1. Reference paper matrix.",
    )

    add_bookmarked_heading(doc, "Appendix B. Commands and Artifacts", "sec_appendix_b", level=1)
    add_docx_paragraph(
        doc,
        "Install and sync with uv sync --extra dev. Build local schema vectors with uv run beacon-index, or uv run python -m beacon.indexing. Load PostgreSQL data with uv run beacon-load-db, or uv run python -m beacon.load_db."
    )
    add_docx_paragraph(
        doc,
        "Run the UI with uv run beacon-ui, or uv run python -m beacon.ui. Run one CLI question with uv run beacon-ask \"...\", or uv run python -m beacon.pipeline \"...\". Run focused tests with uv run pytest tests -v."
    )
    add_docx_paragraph(
        doc,
        "Main artifacts are data/semantic_model/*.json, data/indices/local_vectors/*, data/few_shot_queries.json, tests/test_results/master_plan_evaluation_results.json, tests/test_results/report.html, and optional data/example_candidates.json. Compatibility wrappers such as src/beacon/ui.py, pipeline.py, load_db.py, retrieval.py, and retrieval_tools.py forward old imports and commands to app, runtime, and linking modules.",
    )

    add_bookmarked_heading(doc, "Appendix C. Horizontal Diagram Recreation Prompt", "sec_appendix_c", level=1)
    add_docx_paragraph(
        doc,
        "If the generated figure needs to be recreated in another tool, use this prompt as the design brief."
    )
    add_docx_paragraph(doc, DIAGRAM_RECREATE_PROMPT)

    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleMain",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=23,
            leading=28,
            textColor=colors.HexColor("#1F2937"),
            alignment=TA_LEFT,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Subtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#6B7280"),
            spaceAfter=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Custom",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=16,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2Custom",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyCustom",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=13.5,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Small",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#1F2937"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCEntry",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=13,
            textColor=colors.HexColor("#1F2937"),
            leftIndent=12,
            spaceBefore=2,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TOCPurpose",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=10.5,
            textColor=colors.HexColor("#6B7280"),
            leftIndent=28,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CodeBlock",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=7.8,
            leading=9.4,
            textColor=colors.HexColor("#1F2937"),
            backColor=colors.HexColor("#F8FAFC"),
            borderColor=colors.HexColor("#D9DEE7"),
            borderWidth=0.35,
            borderPadding=5,
            spaceBefore=4,
            spaceAfter=10,
        )
    )
    return styles


def para(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(format_pdf_inline(text), style)


def raw_para(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text, style)


def format_pdf_inline(text: str) -> str:
    parts = []
    for kind, value in split_formatted_tokens(text):
        escaped = html.escape(value)
        if kind == "important":
            parts.append(f"<b>{escaped}</b>")
        elif kind == "path":
            parts.append(f'<font color="#6B7280"><i>{escaped}</i></font>')
        else:
            parts.append(escaped)
    return "".join(parts)


def pdf_link_line(title: str, purpose: str, anchor: str, page: int, styles) -> list:
    safe_title = html.escape(title)
    safe_purpose = html.escape(purpose)
    line = f'<link href="#{anchor}" color="#2E74B5">{safe_title}</link><font color="#6B7280"> . . . . . {page}</font>'
    return [
        raw_para(line, styles["TOCEntry"]),
        raw_para(f'<font color="#6B7280">{safe_purpose}</font>', styles["TOCPurpose"]),
    ]


def pdf_heading(text: str, anchor: str, styles, level: int = 1) -> Paragraph:
    style = styles["H1Custom"] if level == 1 else styles["H2Custom"]
    return raw_para(f'<a name="{anchor}"/>{html.escape(text)}', style)


def pdf_table(headers: list[str], rows: list[tuple[str, ...]], widths: list[float], styles) -> LongTable:
    data = [[para(h, styles["Small"]) for h in headers]]
    for row in rows:
        data.append([para(str(item), styles["Small"]) for item in row])
    table = LongTable(data, colWidths=[w * inch for w in widths], repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9DEE7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def pdf_table_caption(text: str, styles) -> Paragraph:
    return Paragraph(
        text,
        ParagraphStyle(
            "TableCaption",
            parent=styles["Small"],
            textColor=colors.HexColor("#6B7280"),
            spaceBefore=6,
            spaceAfter=4,
        ),
    )


def pdf_figure(path: Path, caption: str, styles, max_width: float = 6.4, max_height: float = 5.2) -> list:
    """Return a scaled PDF image and caption if the figure exists."""
    if not path.exists():
        return []
    with PILImage.open(path) as img:
        width_px, height_px = img.size
    scale = min((max_width * inch) / width_px, (max_height * inch) / height_px)
    figure = Image(str(path), width=width_px * scale, height=height_px * scale)
    caption_style = ParagraphStyle(
        "Caption",
        parent=styles["Small"],
        alignment=TA_CENTER,
        textColor=colors.HexColor("#6B7280"),
        spaceAfter=8,
    )
    return [KeepTogether([figure, Paragraph(caption, caption_style)]), Spacer(1, 6)]


def add_pdf_report() -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        title="Beacon Method and Pipeline Report",
        author="Codex",
    )
    story = []
    story.append(para("Beacon Method and Pipeline Report", styles["TitleMain"]))
    story.append(para("Submission-ready technical method report", styles["Subtitle"]))
    meta_rows = [
        ("Project", "Beacon / basic-mvp"),
        ("Report date", "2026-06-30"),
        ("Template spine", "Oxford Engineering style adapted for a technical method report"),
        ("Scope", "Current MVP plus benchmark-planning notes from the sibling Beacon workspace"),
    ]
    story.append(pdf_table(["Field", "Value"], meta_rows, [1.2, 5.2], styles))
    story.append(Spacer(1, 12))

    story.append(PageBreak())
    story.append(para("Contents", styles["H1Custom"]))
    story.append(
        para(
            "The entries below use static page numbers from the generated PDF and hyperlink directly to the corresponding section.",
            styles["BodyCustom"],
        )
    )
    for title, purpose, anchor, page in CONTENTS_ENTRIES:
        story.extend(pdf_link_line(title, purpose, anchor, page, styles))

    story.append(PageBreak())
    story.append(pdf_heading("1. Executive Summary", "sec_exec", styles))
    story.append(
        para(
            "Beacon is a metadata-grounded natural-language-to-SQL MVP. Its current core pipeline is: profile data, retrieve grounded schema context, build a SQL-only prompt, generate one read-only PostgreSQL query, validate and execute it safely, review the result, retry when useful, and then compose the final answer.",
            styles["BodyCustom"],
        )
    )
    story.append(
        para(
            "The method is intentionally simple: plain dictionaries, small helper modules, and focused tests. The strongest current design idea is grounding: semantic profiles, value evidence, local vector retrieval, join paths, and examples.",
            styles["BodyCustom"],
        )
    )
    story.append(
        para(
            "The main limitation is that selected schema is recall-first and can include distracting context. Reference papers are mapped step-by-step so the method is traceable.",
            styles["BodyCustom"],
        )
    )

    story.append(PageBreak())
    story.append(pdf_heading("2. System Context", "sec_context", styles))
    story.append(
        para(
            "Beacon currently works over a local e-commerce PostgreSQL schema with seven tables: customers, orders, order_items, products, geography, sales, and inventory. Each table has a semantic JSON file under data/semantic_model/. The offline indexing layer enriches those files with compact data evidence instead of dumping raw data into prompts.",
            styles["BodyCustom"],
        )
    )
    story.append(
        para(
            "The semantic layer stores table-level evidence such as semantic name, description, grain, relations, and sample rows. It also stores column-level evidence such as names, types, descriptions, profiles, aliases, and value aliases where useful.",
            styles["BodyCustom"],
        )
    )
    story.append(
        para(
            "Profile fields include null count, distinct count, sample values, numeric/date bounds, means for non-ID numeric fields, and top categorical values. Few-shot examples live in data/few_shot_queries.json and are enriched with structural signals during indexing.",
            styles["BodyCustom"],
        )
    )

    story.append(PageBreak())
    story.append(pdf_heading("3. Current Pipeline", "sec_pipeline", styles))
    story.append(
        para(
            "The current basic-mvp implementation is organized in four layers: indexing, linking, runtime, and app. The important source files are under src/beacon/, with compatibility wrappers kept at the package root.",
            styles["BodyCustom"],
        )
    )
    story.extend(
        pdf_figure(
            CURRENT_PIPELINE_DIAGRAM,
            "Figure 1. Current horizontal Beacon pipeline, generated from the package layout.",
            styles,
            max_width=6.4,
            max_height=4.8,
        )
    )
    story.append(
        para(
            "At runtime, the UI or CLI passes the question into answer_question and answer_section. The section runner retrieves schema context, assembles a SQL-only prompt, requests SQL from the configured model, validates and executes the query through runtime.sql, reviews the result, and either retries with targeted repair or composes the final answer.",
            styles["BodyCustom"],
        )
    )

    story.append(PageBreak())
    story.append(pdf_heading("4. Related Work Positioning", "sec_related", styles))
    story.append(
        para(
            "The reference set was chosen around four practical questions: how to enrich schema context, how to avoid losing required tables, how to organize prompts and examples, and how to recover from SQL errors. Beacon is not trying to reproduce a full competition system such as DAIL-SQL or DIN-SQL. Instead, it extracts the pieces that fit a small, inspectable MVP: semantic metadata, evidence-first linking, stable prompt order, read-only validation, and a bounded retry loop.",
            styles["BodyCustom"],
        )
    )
    story.append(
        para(
            "The strongest thesis is that Beacon improves usefulness through grounding rather than model cleverness. Papers on metadata extraction and question enrichment motivate profile values and matched evidence [P1, P2]. Schema-linking and retrieval papers motivate recall-preserving schema selection with later precision work [P4-P7]. Prompting, retry, and reliability papers motivate the fixed prompt layout, execution feedback, result review, and readable failure states [P8-P14].",
            styles["BodyCustom"],
        )
    )
    story.append(
        para(
            "This positioning also explains the current limitation: the MVP prefers to retrieve enough context, even when that includes extra tables. The benchmark-oriented next step is not a bigger orchestration framework, but a sharper schema-focus layer that preserves recall while reducing distracting columns and projection mismatches.",
            styles["BodyCustom"],
        )
    )

    story.append(PageBreak())
    story.append(pdf_heading("5. Research-to-Method Traceability", "sec_traceability", styles))
    story.append(
        para(
            "This table is the core traceability view. It distinguishes implemented MVP mechanics from planned benchmark extensions.",
            styles["BodyCustom"],
        )
    )
    story.append(pdf_table_caption("Table 2. Research-to-method traceability matrix.", styles))
    story.append(pdf_table(["Pipeline step", "Method source", "Implementation status"], traceability_rows(), [1.35, 3.0, 2.05], styles))

    story.append(PageBreak())
    story.append(pdf_heading("6. Method Details", "sec_method", styles))
    story.append(
        para(
            "This section describes each layer as a method, not just as a file. Reference IDs are cited inline in the same style as a paper; Appendix A expands the titles and benchmark notes.",
            styles["BodyCustom"],
        )
    )
    story.extend(
        pdf_figure(
            METHOD_LAYER_DIAGRAM,
            "Figure 2. Data and artifact flow through Beacon's method layers.",
            styles,
            max_width=6.4,
            max_height=4.8,
        )
    )
    for detail in METHOD_DEEP_DETAILS:
        story.append(para(detail["layer"], styles["H2Custom"]))
        for paragraph in method_detail_paragraphs(detail):
            story.append(para(paragraph, styles["BodyCustom"]))
        if detail["layer"] == "SQL validation and execution":
            story.extend(
                pdf_figure(
                    RETRY_LOOP_DIAGRAM,
                    "Figure 3. SQL validation, execution, review, and retry loop.",
                    styles,
                    max_width=6.4,
                    max_height=3.6,
                )
            )

    story.append(PageBreak())
    story.append(pdf_heading("7. Concrete Example Walkthrough", "sec_walkthrough", styles))
    story.append(
        para(
            "The walkthrough below uses the question 'For paid Apple Pay orders in the East region during 2022, what are the top 5 product categories by net revenue, discount rate, and profit margin?' and records the actual retrieval/linking outputs from a deterministic hash-embedding inspection run. This example is intentionally more demanding: it needs value grounding, date recognition, multi-table joins, grouped metrics, top-k ranking, and a precise SQL output shape.",
            styles["BodyCustom"],
        )
    )
    story.extend(
        pdf_figure(
            WALKTHROUGH_DIAGRAM,
            "Figure 4. Layer-by-layer walkthrough for the complex paid-Apple-Pay regional revenue question.",
            styles,
            max_width=6.1,
            max_height=6.8,
        )
    )
    story.append(pdf_table_caption("Table 3. Observed outputs at each walkthrough layer.", styles))
    story.append(pdf_table(["Layer", "Observed output", "Interpretation"], EXAMPLE_WALKTHROUGH_ROWS, [1.05, 2.95, 2.4], styles))
    story.append(para("Expected SQL Shape For The Example", styles["H2Custom"]))
    story.append(Preformatted(EXPECTED_WALKTHROUGH_SQL, styles["CodeBlock"]))

    story.append(PageBreak())
    story.append(pdf_heading("8. Implementation Call Flow", "sec_call_flow", styles))
    story.append(
        para(
            "This table is intentionally function-level: it shows the call contract after the method layers have already explained the research mechanics.",
            styles["BodyCustom"],
        )
    )
    story.append(pdf_table_caption("Table 4. Function-level call flow.", styles))
    story.append(pdf_table(["Role", "Function or module", "Call contract"], FUNCTION_MAP, [1.25, 2.2, 2.95], styles))

    story.append(PageBreak())
    story.append(pdf_heading("9. Evaluation Status", "sec_evaluation", styles))
    story.append(
        para(
            "Evaluation is still MVP-scale, but it is concrete enough to show the pipeline is testable. The strongest evidence is the combination of focused tests, successful index construction, the latest 10-question rerun, and the benchmark planning note that the best achieved BIRD-dev result is 55% accuracy.",
            styles["BodyCustom"],
        )
    )
    story.append(
        raw_para(
            f'Latest run artifacts folder: <link href="{html.escape(TEST_RESULTS_REPO_PATH)}" color="#2E74B5">{html.escape(TEST_RESULTS_REPO_PATH)}</link>',
            styles["BodyCustom"],
        )
    )
    story.append(pdf_table_caption("Table 5. Current evaluation evidence and caveats.", styles))
    story.append(pdf_table(["Evidence", "Current status", "Notes"], EVALUATION_ROWS, [1.35, 1.45, 3.6], styles))

    story.append(PageBreak())
    story.append(pdf_heading("10. Limitations and Next Work", "sec_limitations", styles))
    story.append(
        para(
            "The next work should follow the same thesis as the method: keep the pipeline grounded and inspectable, then improve precision where the walkthrough exposes extra context.",
            styles["BodyCustom"],
        )
    )
    story.append(pdf_table_caption("Table 6. Prioritized limitations and next actions.", styles))
    story.append(pdf_table(["Priority", "Limitation", "Next work"], LIMITATION_ROWS, [0.7, 2.65, 3.05], styles))

    story.append(PageBreak())
    story.append(pdf_heading("Appendix A. Reference Paper Matrix", "sec_appendix_a", styles))
    story.append(pdf_table_caption("Table A1. Reference paper matrix.", styles))
    story.append(
        pdf_table(
            ["ID", "Paper", "Area", "Benchmark/result", "Beacon use"],
            [(p["id"], p["title"], p["area"], p["benchmark"] + "; " + p["result"], p["beacon_use"]) for p in PAPERS],
            [0.35, 1.7, 0.85, 1.85, 1.65],
            styles,
        )
    )

    story.append(pdf_heading("Appendix B. Commands and Artifacts", "sec_appendix_b", styles))
    story.append(
        para(
            "Install and sync with uv sync --extra dev. Build local schema vectors with uv run beacon-index, or uv run python -m beacon.indexing. Load PostgreSQL data with uv run beacon-load-db, or uv run python -m beacon.load_db.",
            styles["BodyCustom"],
        )
    )
    story.append(
        para(
            "Run the UI with uv run beacon-ui, or uv run python -m beacon.ui. Run one CLI question with uv run beacon-ask \"...\", or uv run python -m beacon.pipeline \"...\". Run focused tests with uv run pytest tests -v.",
            styles["BodyCustom"],
        )
    )
    story.append(
        para(
            "Main artifacts are data/semantic_model/*.json, data/indices/local_vectors/*, data/few_shot_queries.json, tests/test_results/master_plan_evaluation_results.json, tests/test_results/report.html, and optional data/example_candidates.json. Compatibility wrappers such as src/beacon/ui.py, pipeline.py, load_db.py, retrieval.py, and retrieval_tools.py forward old imports and commands to app, runtime, and linking modules.",
            styles["BodyCustom"],
        )
    )
    story.append(pdf_heading("Appendix C. Horizontal Diagram Recreation Prompt", "sec_appendix_c", styles))
    story.append(para(DIAGRAM_RECREATE_PROMPT, styles["BodyCustom"]))
    def footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawString(0.8 * inch, 0.45 * inch, "Beacon Method Report")
        canvas.drawRightString(7.7 * inch, 0.45 * inch, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_visual_assets()
    add_docx_report()
    add_pdf_report()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
